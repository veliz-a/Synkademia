import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generar_docx(proyecto, integrantes, tareas_completadas) -> io.BytesIO:
    """Genera un archivo DOCX en memoria basado en los datos del proyecto."""
    doc = Document()
    
    # Configuración de márgenes estándar (APA)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    metadata = proyecto.project_metadata or {}
    
    # 1. Portada
    doc.add_paragraph()
    doc.add_paragraph()
    
    p_title = doc.add_paragraph(proyecto.title)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.runs[0].font.bold = True
    p_title.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    
    p_course = doc.add_paragraph(f"Curso: {proyecto.course}")
    p_course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if metadata.get("institution"):
        p_inst = doc.add_paragraph(metadata["institution"])
        p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_paragraph()
    p_team = doc.add_paragraph("Autores:")
    p_team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for user in integrantes:
        p_u = doc.add_paragraph(user.username)
        p_u.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 2. Inyección de Contenido
    for tarea in tareas_completadas:
        if tarea.content and tarea.content.strip():
            # Título de la sección
            h = doc.add_heading(tarea.title, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Procesamiento básico de marcadores y texto
            parrafos = str(tarea.content).split('\n\n')
            for parrafo in parrafos:
                texto_limpio = parrafo.strip()
                if not texto_limpio:
                    continue
                    
                # Manejo simple de los marcadores del editor
                if texto_limpio.startswith("[ESTRUCTURA_TABLA"):
                    p = doc.add_paragraph("<< Marcador de Tabla detectado >>")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif texto_limpio.startswith("[ESTRUCTURA_IMAGEN"):
                    p = doc.add_paragraph("<< Marcador de Imagen detectado >>")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p = doc.add_paragraph(texto_limpio)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    # Sangría de primera línea (Estilo APA)
                    p.paragraph_format.first_line_indent = Inches(0.5)

    # 3. Guardar en buffer de memoria
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer